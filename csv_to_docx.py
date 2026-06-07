import pandas as pd
from docx import Document
from pathlib import Path

out_dir = Path('outputs')
out_dir.mkdir(exist_ok=True)

csv_path = out_dir / 'descriptive_stats.csv'
if not csv_path.exists():
    raise FileNotFoundError("outputs/descriptive_stats.csv not found. Run descriptive_stats.py first.")

df = pd.read_csv(csv_path, index_col=0)

doc = Document()
doc.add_heading('Descriptive statistics', level=1)

rows, cols = df.shape
table = doc.add_table(rows=rows + 1, cols=cols + 1)
table.style = 'Table Grid'

# Header row
hdr_cells = table.rows[0].cells
hdr_cells[0].text = ''
for j, col in enumerate(df.columns, start=1):
    hdr_cells[j].text = str(col)

# Data rows
for i, idx in enumerate(df.index, start=1):
    row_cells = table.rows[i].cells
    row_cells[0].text = str(idx)
    for j, col in enumerate(df.columns, start=1):
        val = df.iloc[i - 1, j - 1]
        row_cells[j].text = str(val)

out_path = out_dir / 'descriptive_stats.docx'
doc.save(out_path)
print(f"Таблиця збережена у файл {out_path}")
