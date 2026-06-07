from pathlib import Path
from docx import Document
from datetime import datetime

out_dir = Path('outputs')
out_dir.mkdir(exist_ok=True)

src = out_dir / 'descriptive_stats.docx'
if not src.exists():
    raise FileNotFoundError('outputs/descriptive_stats.docx not found. Run csv_to_docx.py first.')

doc = Document(src)

# Додаємо підпис та дату
now = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
p = doc.add_paragraph()
p.add_run('\nПідпис: ').bold = True
p.add_run('згенеровано скриптом').italic = True
p = doc.add_paragraph()
p.add_run('Дата генерації: ').bold = True
p.add_run(now)

out = out_dir / 'descriptive_stats_signed.docx'
doc.save(out)
print(f"Signed document saved: {out.resolve()}")
